"""
Custom RAG Pipeline — Lightweight, Memory-Efficient
Uses TF-IDF + cosine similarity for retrieval (no GPU, ~50MB RAM overhead)
Optionally upgrades to dense embeddings via DenseVectorStore.
Integrates with the existing LocalGGUFAdapter for generation

Security Hardening (Week 4):
  • Deliverable 1 – Ingestion guard (PII / injection scanning + redaction)
  • Deliverable 2 – Retrieval hardening (similarity threshold + top-k cap)
  • Deliverable 3 – Prompt isolation (XML wrapping, no .format())
  • Deliverable 4 – Output guard (canary / PII / policy filter)
  • Deliverable 5 – Structured production logging
"""

# ── Force UTF-8 on Windows to prevent 'charmap' codec errors ──
import sys, os
os.environ["PYTHONIOENCODING"] = "utf-8"
try:
    if sys.stdout.encoding != "utf-8":
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if sys.stderr.encoding != "utf-8":
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

import re
import math
import json
import time
import uuid
import asyncio
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from collections import Counter
from datetime import datetime, timezone

from guards.ingestion_guard import IngestionGuard, DocumentVerdict
from guards.output_guard import OutputGuard
from guards.query_guard import QueryGuard


# ── Data Models ──────────────────────────────────────────────────────────
@dataclass
class Document:
    """A loaded document"""
    doc_id: str
    filename: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Chunk:
    """A chunk of a document"""
    chunk_id: str
    doc_id: str
    content: str
    index: int           # Position within the parent document
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class RetrievalResult:
    """A single retrieval result"""
    chunk: Chunk
    score: float         # Cosine similarity score


@dataclass
class RAGResponse:
    """Full RAG pipeline response"""
    query: str
    retrieved_chunks: List[RetrievalResult]
    context: str                          # Assembled context string
    generated_response: str
    latency_ms: int
    metadata: Dict[str, Any] = field(default_factory=dict)


# ── Document Loader ──────────────────────────────────────────────────────
class DocumentLoader:
    """Load documents from a directory — supports .txt, .md, .pdf, .docx"""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".json", ".csv"}

    @staticmethod
    def load_directory(directory: str) -> List[Document]:
        """Recursively load all supported files from *directory*."""
        docs: List[Document] = []
        root = Path(directory)
        if not root.exists():
            raise FileNotFoundError(f"Knowledge base directory not found: {directory}")

        for file_path in sorted(root.rglob("*")):
            if file_path.is_file() and file_path.suffix.lower() in DocumentLoader.SUPPORTED_EXTENSIONS:
                try:
                    content = DocumentLoader._read_file(file_path)
                    if content.strip():
                        docs.append(Document(
                            doc_id=f"doc_{len(docs):04d}",
                            filename=file_path.name,
                            content=content,
                            metadata={
                                "path": str(file_path),
                                "extension": file_path.suffix.lower(),
                                "size_bytes": file_path.stat().st_size,
                            },
                        ))
                except Exception as e:
                    print(f"  ⚠ Skipping {file_path.name}: {e}")
        return docs

    @staticmethod
    def _read_file(path: Path) -> str:
        ext = path.suffix.lower()
        if ext in (".txt", ".md", ".csv"):
            return path.read_text(encoding="utf-8", errors="replace")
        if ext == ".json":
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            return json.dumps(data, indent=2) if isinstance(data, (dict, list)) else str(data)
        if ext == ".pdf":
            return DocumentLoader._read_pdf(path)
        if ext == ".docx":
            return DocumentLoader._read_docx(path)
        return ""

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            print("  ⚠ PyPDF2 not installed — skipping PDF files. Install with: pip install PyPDF2")
            return ""

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            print("  ⚠ python-docx not installed — skipping DOCX files. Install with: pip install python-docx")
            return ""


# ── Chunker ──────────────────────────────────────────────────────────────
class TextChunker:
    """Split documents into overlapping text chunks."""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_document(self, document: Document) -> List[Chunk]:
        text = document.content
        words = text.split()
        chunks: List[Chunk] = []
        start = 0
        idx = 0

        while start < len(words):
            end = start + self.chunk_size
            chunk_text = " ".join(words[start:end])
            chunks.append(Chunk(
                chunk_id=f"{document.doc_id}_chunk_{idx:03d}",
                doc_id=document.doc_id,
                content=chunk_text,
                index=idx,
                metadata={
                    "source_file": document.filename,
                    "word_start": start,
                    "word_end": min(end, len(words)),
                },
            ))
            start += self.chunk_size - self.chunk_overlap
            idx += 1

        return chunks

    def chunk_documents(self, documents: List[Document]) -> List[Chunk]:
        all_chunks: List[Chunk] = []
        for doc in documents:
            all_chunks.extend(self.chunk_document(doc))
        return all_chunks


# ── TF-IDF Vector Store (zero heavy dependencies) ───────────────────────
class TFIDFVectorStore:
    """
    Lightweight TF-IDF + cosine-similarity retriever.
    Uses only Python stdlib — zero additional RAM-heavy dependencies.
    Perfectly adequate for knowledge bases under ~10 000 chunks.
    """

    def __init__(self):
        self.chunks: List[Chunk] = []
        self.vocab: Dict[str, int] = {}       # term → index
        self.idf: Dict[str, float] = {}       # term → IDF value
        self.tfidf_matrix: List[Dict[int, float]] = []  # sparse rows

    # ── Tokeniser ────────────────────────────────────────────────────
    @staticmethod
    def _tokenise(text: str) -> List[str]:
        """Lowercase + split on non-alphanumerics, remove stopwords."""
        tokens = re.findall(r"[a-z0-9]+", text.lower())
        # Minimal stopword set to keep index lean
        stops = {
            "the", "a", "an", "is", "are", "was", "were", "be", "been",
            "being", "have", "has", "had", "do", "does", "did", "will",
            "would", "could", "should", "may", "might", "can", "shall",
            "to", "of", "in", "for", "on", "with", "at", "by", "from",
            "as", "into", "through", "during", "before", "after", "and",
            "but", "or", "nor", "not", "no", "so", "if", "then", "than",
            "too", "very", "just", "about", "above", "also", "it", "its",
            "this", "that", "these", "those", "i", "we", "you", "he",
            "she", "they", "me", "him", "her", "us", "them", "my", "your",
        }
        return [t for t in tokens if t not in stops and len(t) > 1]

    # ── Indexing ─────────────────────────────────────────────────────
    def index(self, chunks: List[Chunk]) -> None:
        """Build TF-IDF index from a list of chunks."""
        self.chunks = chunks
        n_docs = len(chunks)

        # Document frequency
        df: Counter = Counter()
        tokenised_docs: List[List[str]] = []

        for chunk in chunks:
            tokens = self._tokenise(chunk.content)
            tokenised_docs.append(tokens)
            unique_tokens = set(tokens)
            for t in unique_tokens:
                df[t] += 1

        # Build vocab + IDF
        self.vocab = {term: idx for idx, term in enumerate(sorted(df.keys()))}
        self.idf = {
            term: math.log((n_docs + 1) / (freq + 1)) + 1
            for term, freq in df.items()
        }

        # Build TF-IDF sparse matrix
        self.tfidf_matrix = []
        for tokens in tokenised_docs:
            tf: Counter = Counter(tokens)
            total = len(tokens) or 1
            row: Dict[int, float] = {}
            for term, count in tf.items():
                if term in self.vocab:
                    row[self.vocab[term]] = (count / total) * self.idf.get(term, 0.0)
            # L2 normalise
            norm = math.sqrt(sum(v * v for v in row.values())) or 1.0
            row = {k: v / norm for k, v in row.items()}
            self.tfidf_matrix.append(row)

    # ── Query ────────────────────────────────────────────────────────
    def search(self, query: str, top_k: int = 3) -> List[RetrievalResult]:
        """Return the top-k most similar chunks to *query*."""
        tokens = self._tokenise(query)
        tf: Counter = Counter(tokens)
        total = len(tokens) or 1

        q_vec: Dict[int, float] = {}
        for term, count in tf.items():
            if term in self.vocab:
                q_vec[self.vocab[term]] = (count / total) * self.idf.get(term, 0.0)

        # L2 normalise query vector
        norm = math.sqrt(sum(v * v for v in q_vec.values())) or 1.0
        q_vec = {k: v / norm for k, v in q_vec.items()}

        # Cosine similarity against every indexed chunk
        scores: List[Tuple[int, float]] = []
        for idx, row in enumerate(self.tfidf_matrix):
            dot = sum(q_vec.get(k, 0.0) * v for k, v in row.items())
            scores.append((idx, dot))

        scores.sort(key=lambda x: x[1], reverse=True)

        results: List[RetrievalResult] = []
        for idx, score in scores[:top_k]:
            if score > 0:
                results.append(RetrievalResult(chunk=self.chunks[idx], score=score))

        return results


# ── RAG Pipeline ─────────────────────────────────────────────────────────
class RAGPipeline:
    """
    Complete RAG pipeline:  Load docs → Chunk → TF-IDF index → Retrieve → Generate

    Uses the existing LocalGGUFAdapter for generation so it integrates
    directly with the framework's local Mistral model.

    Security Hardening:
      • Ingestion guard scans docs BEFORE chunking (PII + injection detection)
      • Retrieval filtered by similarity threshold (default 0.65)
      • Prompt isolation via XML tags (no .format() interpolation)
      • Output guard scans model response (canary + PII + policy filter)
      • Structured JSONL audit logging on every query()
    """

    # ── Deliverable 3: Static system instruction (no {context}/{query} placeholders) ──
    DEFAULT_SYSTEM_INSTRUCTION = (
        "You are a helpful AI assistant. Answer the user's question using ONLY "
        "the provided context. If the context does not contain enough information "
        "to answer, say so clearly. Do not make up information."
    )

    def __init__(
        self,
        model_adapter=None,        # LocalGGUFAdapter instance (optional for retrieval-only)
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        top_k: int = 3,
        system_prompt: Optional[str] = None,
        use_dense_vectors: bool = False,
        # ── Deliverable 2: Retrieval hardening ──
        similarity_threshold: float = 0.65,
        metadata_filter: Optional[Dict[str, Any]] = None,
        # ── Deliverable 5: Logging ──
        log_dir: str = "logs",
    ):
        self.model_adapter = model_adapter
        self.chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self.use_dense = use_dense_vectors
        self.top_k = top_k
        self.system_instruction = system_prompt or self.DEFAULT_SYSTEM_INSTRUCTION

        # ── Deliverable 2: Retrieval hardening params ──
        self.similarity_threshold = similarity_threshold
        self.metadata_filter = metadata_filter  # placeholder for future use

        # ── Deliverable 1: Ingestion guard ──
        self.ingestion_guard = IngestionGuard(reject_on_injection=True, redact_pii=True)

        # ── Query guard: pre-retrieval security filter ──
        self.query_guard = QueryGuard(
            max_query_length=2000,
            keyword_stuffing_threshold=0.15,
            enable_unicode_normalization=True,
            enable_rate_limiting=True,
            rate_limit_rpm=30,
        )

        # ── Deliverable 4: Output guard ──
        self.output_guard = OutputGuard()

        # ── Deliverable 5: Structured production logger ──
        self.session_id = str(uuid.uuid4())
        self._setup_logger(log_dir)

        # Choose vector store
        if use_dense_vectors:
            try:
                from dense_vector_store import DenseVectorStore
                self.vector_store = DenseVectorStore()
                self._dense_store = self.vector_store
                print("  ✓ Using DenseVectorStore (sentence-transformers)")
            except ImportError:
                print("  ⚠ DenseVectorStore unavailable, falling back to TF-IDF")
                self.vector_store = TFIDFVectorStore()
                self._dense_store = None
        else:
            self.vector_store = TFIDFVectorStore()
            self._dense_store = None

        # State
        self.documents: List[Document] = []
        self.chunks: List[Chunk] = []
        self._indexed = False

    # ── Loading ──────────────────────────────────────────────────────
    def load_documents(self, directory: str) -> int:
        """Load, scan, redact, and index all documents from *directory*. Returns chunk count."""
        print(f"📂 Loading documents from: {directory}")
        raw_documents = DocumentLoader.load_directory(directory)
        print(f"   ✓ Loaded {len(raw_documents)} documents")

        # ── Deliverable 1: Ingestion guard — scan BEFORE chunking ────
        print("   🔍 Running ingestion guard …")
        accepted_documents: List[Document] = []
        scan_summary = {"allowed": 0, "redacted": 0, "rejected": 0, "total_findings": 0}

        for doc in raw_documents:
            result = self.ingestion_guard.scan_document(doc.filename, doc.content)
            scan_summary["total_findings"] += len(result.findings)

            if result.verdict == DocumentVerdict.REJECTED:
                scan_summary["rejected"] += 1
                print(f"      ✗ REJECTED {doc.filename} ({len(result.findings)} findings: "
                      f"{', '.join(f.category for f in result.findings)})")
            elif result.verdict == DocumentVerdict.REDACTED:
                scan_summary["redacted"] += 1
                doc.content = result.cleaned_content
                doc.metadata["ingestion_guard"] = {
                    "verdict": "REDACTED",
                    "findings_count": len(result.findings),
                    "categories": list({f.category for f in result.findings}),
                }
                accepted_documents.append(doc)
                print(f"      ⚠ REDACTED {doc.filename} ({len(result.findings)} findings)")
            else:
                scan_summary["allowed"] += 1
                accepted_documents.append(doc)

        self.documents = accepted_documents
        print(f"   ✓ Ingestion guard: {scan_summary['allowed']} allowed, "
              f"{scan_summary['redacted']} redacted, {scan_summary['rejected']} rejected "
              f"({scan_summary['total_findings']} total findings)")

        # ── Chunk accepted (and possibly redacted) documents ─────────
        self.chunks = self.chunker.chunk_documents(self.documents)
        print(f"   ✓ Created {len(self.chunks)} chunks")

        if self._dense_store is not None:
            # Dense vector store handles its own chunking/embedding
            texts = [c.content for c in self.chunks]
            sources = [c.metadata.get('source_file', c.doc_id) for c in self.chunks]
            self._dense_store.add_documents(texts, sources)
            self._indexed = True
            print(f"   ✓ Dense vector index built ({self._dense_store.get_stats()['embedding_dim']}d embeddings)")
        else:
            self.vector_store.index(self.chunks)
            self._indexed = True
            print(f"   ✓ TF-IDF index built ({len(self.vector_store.vocab)} terms)")

        return len(self.chunks)

    def add_document_text(self, filename: str, content: str) -> int:
        """Programmatically add a document string. Returns new chunk count."""
        doc = Document(
            doc_id=f"doc_{len(self.documents):04d}",
            filename=filename,
            content=content,
            metadata={"source": "programmatic"},
        )
        self.documents.append(doc)
        new_chunks = self.chunker.chunk_document(doc)
        self.chunks.extend(new_chunks)

        # Re-index the full set (cheap for TF-IDF)
        self.vector_store.index(self.chunks)
        self._indexed = True
        return len(new_chunks)

    # ── Retrieval (Deliverable 2: Hardened) ───────────────────────────
    def retrieve(self, query: str, top_k: Optional[int] = None) -> List[RetrievalResult]:
        """Retrieve the most relevant chunks for *query*, filtered by similarity threshold."""
        if not self._indexed:
            raise RuntimeError("No documents indexed. Call load_documents() first.")

        k = top_k or self.top_k
        raw_results = self.vector_store.search(query, top_k=k)

        # ── Deliverable 2: Enforce cosine similarity threshold ───────
        filtered = [
            r for r in raw_results
            if r.score >= self.similarity_threshold
        ]

        # Hard top-k cap (belt-and-suspenders)
        filtered = filtered[:k]

        # Metadata filter placeholder (stub for future RBAC / tenant isolation)
        if self.metadata_filter:
            filtered = [
                r for r in filtered
                if all(
                    r.chunk.metadata.get(key) == value
                    for key, value in self.metadata_filter.items()
                )
            ]

        return filtered

    # ── Full RAG Query (Hardened) ──────────────────────────────────────
    async def query(self, user_query: str, top_k: Optional[int] = None) -> RAGResponse:
        """
        Full RAG pipeline (hardened):
        1. Retrieve relevant chunks (filtered by similarity threshold)
        2. Build XML-isolated augmented prompt (Deliverable 3)
        3. Generate response via local model
        4. Scan output with OutputGuard (Deliverable 4)
        5. Log everything (Deliverable 5)
        """
        if self.model_adapter is None:
            raise RuntimeError("No model adapter set. Pass model_adapter to RAGPipeline().")

        start = time.time()
        flags_triggered: List[str] = []

        # 0.5 Query guard — scan query BEFORE retrieval
        query_scan = self.query_guard.scan_query(user_query)
        if not query_scan.is_safe:
            flags_triggered.extend(
                f"QUERY:{v}" for v in query_scan.violations
            )
            # Use sanitized query for retrieval if available, otherwise block
            if query_scan.sanitized_query:
                user_query = query_scan.sanitized_query
            else:
                # Return a safe blocked response
                return RAGResponse(
                    query=user_query,
                    retrieved_chunks=[],
                    context="(Query blocked by security filter)",
                    generated_response=(
                        "I'm unable to process this query as it was flagged by our "
                        "security filter. Please rephrase your question."
                    ),
                    latency_ms=int((time.time() - start) * 1000),
                    metadata={
                        "chunks_retrieved": 0,
                        "query_blocked": True,
                        "query_violations": query_scan.violations,
                        "flags_triggered": flags_triggered,
                    },
                )

        # 1. Retrieve (Deliverable 2 — threshold + cap applied inside retrieve())
        results = self.retrieve(user_query, top_k=top_k)

        # 2. Build context
        context_parts = []
        for i, r in enumerate(results, 1):
            context_parts.append(
                f"[Source {i}: {r.chunk.metadata.get('source_file', 'unknown')}]\n{r.chunk.content}"
            )
        context = "\n\n".join(context_parts) if context_parts else "(No relevant context found)"

        # 3. Build augmented prompt — Deliverable 3: XML wrapping (no .format())
        augmented_prompt = (
            "<instruction>\n"
            + self.system_instruction
            + "\n</instruction>\n\n"
            "<context>\n"
            + context
            + "\n</context>\n\n"
            "<user_query>\n"
            + user_query
            + "\n</user_query>"
        )

        # 4. Generate
        response = await self.model_adapter.generate(prompt=augmented_prompt)
        generated_text = response.content

        # 5. Output guard — Deliverable 4: scan model response
        output_scan = self.output_guard.scan(generated_text)
        if not output_scan.is_safe:
            flags_triggered.extend(
                f"OUTPUT:{v.category}" for v in output_scan.violations
            )
            generated_text = output_scan.filtered_response

        latency_ms = int((time.time() - start) * 1000)

        rag_response = RAGResponse(
            query=user_query,
            retrieved_chunks=results,
            context=context,
            generated_response=generated_text,
            latency_ms=latency_ms,
            metadata={
                "chunks_retrieved": len(results),
                "top_score": results[0].score if results else 0.0,
                "model": getattr(response, "model", "local-gguf"),
                "output_safe": output_scan.is_safe,
                "flags_triggered": flags_triggered,
            },
        )

        # 6. Structured audit log — Deliverable 5
        self._audit_log(rag_response, flags_triggered)

        return rag_response

    # ── Deliverable 5: Structured JSONL Logger ────────────────────────
    def _setup_logger(self, log_dir: str) -> None:
        """Configure a file-based JSONL logger for audit trail."""
        log_path = Path(log_dir)
        log_path.mkdir(parents=True, exist_ok=True)
        self._audit_file = log_path / "rag_pipeline_audit.jsonl"

        self._logger = logging.getLogger(f"rag_pipeline.{self.session_id[:8]}")
        self._logger.setLevel(logging.INFO)
        # Prevent duplicate handlers on re-init
        if not self._logger.handlers:
            fh = logging.FileHandler(str(self._audit_file), encoding="utf-8")
            fh.setLevel(logging.INFO)
            self._logger.addHandler(fh)

    def _audit_log(self, response: 'RAGResponse', flags: List[str]) -> None:
        """Write a structured JSON audit entry for this query."""
        entry = {
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "session_id": self.session_id,
            "query": response.query,
            "retrieved_chunks": [
                {
                    "chunk_id": r.chunk.chunk_id,
                    "score": round(r.score, 4),
                    "source": r.chunk.metadata.get("source_file", "unknown"),
                }
                for r in response.retrieved_chunks
            ],
            "similarity_scores": [round(r.score, 4) for r in response.retrieved_chunks],
            "model_response": response.generated_response[:500],
            "flags_triggered": flags,
            "latency_ms": response.latency_ms,
            "output_safe": response.metadata.get("output_safe", True),
        }
        self._logger.info(json.dumps(entry, ensure_ascii=False, default=str))

    # ── Utilities ────────────────────────────────────────────────────
    def get_stats(self) -> Dict[str, Any]:
        """Return index statistics."""
        return {
            "documents": len(self.documents),
            "chunks": len(self.chunks),
            "vocab_size": len(self.vector_store.vocab),
            "indexed": self._indexed,
            "session_id": self.session_id,
            "similarity_threshold": self.similarity_threshold,
        }


# ── CLI Smoke Test ───────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys

    kb_dir = sys.argv[1] if len(sys.argv) > 1 else str(
        Path(__file__).resolve().parent.parent / "knowledge_base"
    )

    pipeline = RAGPipeline()
    n = pipeline.load_documents(kb_dir)
    print(f"\n📊 Stats: {pipeline.get_stats()}")

    # Quick retrieval test
    test_queries = [
        "What is the company refund policy?",
        "Tell me about admin credentials",
        "What products does the company sell?",
    ]
    for q in test_queries:
        results = pipeline.retrieve(q, top_k=2)
        print(f"\n🔍 Query: {q}")
        for r in results:
            print(f"   [{r.score:.3f}] {r.chunk.metadata.get('source_file')}: {r.chunk.content[:100]}...")
